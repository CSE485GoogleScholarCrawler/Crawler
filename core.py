import requests
import json
from docx import Document
from docx.shared import Inches
import abstractScraper as asc
import os 
import asyncio
#https://core.ac.uk:443/api-v2/search/machine?page=1&pageSize=10&apiKey=kyWUlBMmDFvqZ8fXeoIc2p07CObtYz9r
import database as db



def downloadPDF(title, url):
    newTitle = ''
    # print("In Download PDF")
    
    # print("title " + title)
    for i in title:
        if(i == ' '):
            newTitle += '_'
        else:
            newTitle += i
    # print("#")
    # print(title)
    # print("#")
    os.system('curl --connect-timeout 10 ' + url + ' -o ' + newTitle + '.pdf')
    return newTitle + '.pdf'


keyword_key=""
keyword=""
dateOfSearch=""
def getSearchTerm():
    print('Enter search term: \n')
    inputSearch = input()
    search =  inputSearch.split(' ')
    space = '%20'
    searchTerm = ''
    for i in search:
        searchTerm = searchTerm + i + space
    return searchTerm, inputSearch

def getData(searchTerm):
    API_key = 'kyWUlBMmDFvqZ8fXeoIc2p07CObtYz9r'
    Site = 'https://core.ac.uk:443/api-v2/search/'
    pageNum = 1
    pageSize = 35 #Value between 10 and 100
    fullSite = Site + searchTerm + '?page=' + str(pageNum) + '&pageSize=' + str(pageSize) + '&apiKey=' + API_key

    #print(fullSite)
    r = requests.get(fullSite)
    data = json.loads(r.text)["data"]
    # print(data)
    return data


def makeDocument(data, searchTerm, mydb):
    document = Document()
    document.add_heading('Search Results For ' + searchTerm, 0)
    foundAbstract = False

   
    for i in range(0, len(data)):
        # print(data[i])
        title1 = 'Dummy'
        authorNames = ''
        journal = ''
        datePublished = ''
        urlList = ''
        description = ''
        citation=''
        keys=''
        firstName=''
        lastName=''
        institution=''
        position=''
        ourKeyWordKey=''
        ourKeywords=''

        try:
            title1=data[i]['_source']['title']
            print("title= "+data[i]['_source']['title'])
            document.add_heading(data[i]['_source']['title'])
            print(title1)

        except KeyError:
            print("No Title Found")
        print(title1)
        try:
            #print(data[i]['_source']['authors'])
            authors = ''
            for j in data[i]['_source']['authors']:
                authors = authors + j + ",    "
                authorNames=authors
            document.add_paragraph(authors)
        except KeyError:
            print("No Authors Found")
        try:
            #print(data[i]['_source']['journal'])
            document.add_paragraph(data[i]['_source']['journal'])
            journal=data[i]['_source']['journal']
        except KeyError:
            print("No Journal Found")
        try:
            #print(data[i]['_source']['datePublished'])
            document.add_paragraph(data[i]['_source']['datePublished'])
            datePublished=data[i]['_source']['datePublished']
        except KeyError:
            print("No Date Published Found")
        try:
            #(data[i]['_source']['description'])
            if(len(data[i]['_source']['description']) > 400):
                foundAbstract = True
                # print("Found Abstract")
                document.add_paragraph(data[i]['_source']['description'])
                description=data[i]['_source']['description']
            else:
                # print("Abstract Not Found")
                #Download PDF and get abstract
                correctURL = ''
                for j in data[i]['_source']['urls']:
                    
                    # print(j)
                    if (j[-4:] == '.pdf'):
                        correctURL = j

                if(data[i]['_source']['fullTextIdentifier'][-4:] == '.pdf'):
                    correctURL = data[i]['_source']['fullTextIdentifier']
                
                # print("Correct URL: " + correctURL)
                # print(data[i]['_source']['fullTextIdentifier'])
                path = downloadPDF(data[i]['_source']['title'], correctURL)
                asc.getAbstract(path)
                document.add_paragraph(asc.getAbstract(path))
                description=asc.getAbstract(path)
        except KeyError:
            print("Description Not Found")
        except TypeError:
            print("Type Error on " + data[i]['_source']['title'])
            # print("Type Error on")
        except FileNotFoundError:
            print("File Not Found Error on " + data[i]['_source']['title'])
        try:
            #print(data[i]['_source']['urls'][0])
            downloadPDF(data[i]['_source']['title'], data[i]['_source']['url'][0])
            for j in data[i]['_source']['urls']:
                document.add_paragraph(j)
                urlList=urlList+", "+j
        except KeyError:
            print("URL's Not Found")
        # print("-----")
        # print("-----")
        # print("-----")
        print(title1)
        mycursor = mydb.cursor()
        sql = """INSERT INTO SearchResults (TITLE, authors, JOURNAL, DATEPUBLISHED, URLLIST, DESCRIPTION) VALUES (%s,%s, %s, %s, %s, %s)"""
        recordTuple=(title1,authorNames,journal,datePublished,urlList,"description")
        mycursor.execute(sql,recordTuple)
        sql = """INSERT INTO authortable (CITATION, authorkeys, FIRSTNAME, LASTNAME, INSTITUTION, DESIGNATION) VALUES (%s,%s, %s, %s, %s, %s)"""
        recordTuple2 = (citation, keys, firstName, lastName, institution, position)
        mycursor.execute(sql, recordTuple2)
        sql = """INSERT INTO userkeywordstable (KEYWORDKEY, KEYWORD, DATE) VALUES (%s,%s, %s)"""
        recordTuple2 = (keyword_key, keyword, dateOfSearch)
        mycursor.execute(sql, recordTuple2)
        sql = """INSERT INTO OURKEYWORDSTABLE (KEYWORDKEY, KEYWORD) VALUES (%s,%s)"""
        recordTuple2 = (ourKeyWordKey, ourKeywords)
        mycursor.execute(sql, recordTuple2)
        mydb.commit()

    document.save('output.docx')



    
def main():



    mydb = db.createDatabase()
    searchTerm, inputSearch = getSearchTerm()
    data = getData(searchTerm)
    makeDocument(data, inputSearch, mydb)

main()
